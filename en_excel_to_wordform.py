from openpyxl import load_workbook
from docx import Document
from itertools import islice
from docx.shared import Pt

# Variables used
fname, fsize = 'Arial', Pt(10) # Font and font size of characters inserted into the form
template_path = "C:/excel_to_wordform/form.docx"  # Path to the form template made as a Word file
excel_path = "C:/excel_to_wordform/data.xlsx"  # Path to the Excel file containing the data to fill out the form. Data for filling out one form is contained in one line
output_folder = "C:/excel_to_wordform/out"  # Folder where the generated Word documents will be saved

def analyze_template(template_path):
    # Open the Word form template
    doc = Document(template_path)

    table_info = []
    
    # Go through all the tables in the template
    for table_index, table in enumerate(doc.tables):
        # Remember the table number and the number of rows and columns
        table_info.append({
            'table_index': table_index,
            'rows': len(table.rows),
            'columns': len(table.columns)
        })
    
    return table_info

def fill_word_template(template_path, excel_path, output_folder):
    # Analyze the template and get information about the tables
    table_info = analyze_template(template_path)
    
    # Loading the Excel file with data to fill
    wb = load_workbook(excel_path)
    ws = wb.active
         
    # Go through the Excel rows, starting with the second row, since the first row usually contains headers
    for row in ws.iter_rows(min_row=2, values_only=True):
        # Get a tuple of length 100 values or less
        row_values = tuple(islice(row, 100))
        # Fill in missing values with empty lines
        row_values += ('',) * (100 - len(row_values))
        
        # Open the Word form template
        doc = Document(template_path)
        # Filling text fields in a Word document with data from an Excel row
        for i, cell_value in enumerate(row):
            # Replace "None" with empty lines
            if cell_value is None:
                cell_value = ''
            field_name = f"{{{i}}}"  # It assumes that fields in a Word document have numbers in curly braces (for example, {0}, {1}, {2}, etc.)
            for paragraph in doc.paragraphs:
                if field_name in paragraph.text:
                    paragraph.text = paragraph.text.replace(field_name, str(cell_value))
                    for run in paragraph.runs:
                        run.font.name = fname  # Setting font of characters inserted into the form
                        run.font.size = fsize    # Setting font size of characters inserted into the form

        # Filling text fields in Word document tables with data from an Excel line
        for table_data in table_info:
            table_index = table_data['table_index']
            rows = table_data['rows']
            columns = table_data['columns']

            # Getting access to the current table
            table = doc.tables[table_index]
            
            # Filling table cells with data from an Excel row
            for i in range(rows):
                for j in range(columns):
                    # Getting cell value from Excel row
                    if i * columns + j < len(row_values):
                        cell_value = row_values[i * columns + j]  # It is assumed that data from Excel is placed into the table row by row
                    else:
                        cell_value = ""
                    cell = table.cell(i, j)

                    # Go through all the paragraphs in the cell and replace the text fields with values ​​from Excel
                    for paragraph in cell.paragraphs:
                        for field_index, field_value in enumerate(row_values):
                            field_name = f"{{{field_index}}}"
                            if field_name in paragraph.text:
                                # Check if the value is "None" and replace it with an empty string
                                if field_value is None:
                                    field_value = ''
                                paragraph.text = paragraph.text.replace(field_name, str(field_value))
                                for run in paragraph.runs:
                                    run.font.name = fname  # Setting font of characters inserted into the form
                                    run.font.size = fsize    # Setting font size of characters inserted into the form

        # Create a name for the new Word file
        output_file = f"{output_folder}/{row[0]}.docx"  # Let's assume that the first Excel column contains unique identifiers

        # Save a new Word document
        doc.save(output_file)

fill_word_template(template_path, excel_path, output_folder)
