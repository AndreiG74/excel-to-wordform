from openpyxl import load_workbook
from docx import Document
from itertools import islice
from docx.shared import Pt

#переменные, использованные в скрипте
fname, fsize = 'Arial', Pt(10) # шрифт
template_path = "C:/excel_to_wordform/form.docx"  # Путь к шаблону Word
excel_path = "C:/excel_to_wordform/data.xlsx"  # Путь к файлу Excel с данными
output_folder = "C:/excel_to_wordform/out"  # Папка, в которой будут сохранены сгенерированные документы Word

def analyze_template(template_path):
    # Открываем шаблон Word
    doc = Document(template_path)

    table_info = []
    
    # Проходим по всем таблицам в шаблоне
    for table_index, table in enumerate(doc.tables):
        # Запоминаем номер таблицы и количество строк и столбцов
        table_info.append({
            'table_index': table_index,
            'rows': len(table.rows),
            'columns': len(table.columns)
        })
    
    return table_info

def fill_word_template(template_path, excel_path, output_folder):
    # Анализируем шаблон и получаем информацию о таблицах
    table_info = analyze_template(template_path)
    
    # Загружаем файл Excel
    wb = load_workbook(excel_path)
    ws = wb.active
         
    # Проходим по строкам Excel, начиная со второй строки, так как первая строка обычно содержит заголовки
    for row in ws.iter_rows(min_row=2, values_only=True):
        # Получаем кортеж длиной 100 значений или меньше
        row_values = tuple(islice(row, 100))
        # Заполняем недостающие значения пустыми строками
        row_values += ('',) * (100 - len(row_values))
        
        # Открываем шаблон Word
        doc = Document(template_path)
        # Заполняем текстовые поля в документе Word данными из строки Excel
        for i, cell_value in enumerate(row):
            # Заменяем "None" на пустые строки
            if cell_value is None:
                cell_value = ''
            field_name = f"{{{i}}}"  # Предполагается, что поля в документе Word имеют номера в фигурных скобках (например, {0}, {1}, {2}, и т. д.)
            for paragraph in doc.paragraphs:
                if field_name in paragraph.text:
                    paragraph.text = paragraph.text.replace(field_name, str(cell_value))
                    for run in paragraph.runs:
                        run.font.name = fname  # устанавливаем шрифт
                        run.font.size = fsize    # устанавливаем размер текста

        # Заполняем таблицы в документе Word данными из строки Excel
        for table_data in table_info:
            table_index = table_data['table_index']
            rows = table_data['rows']
            columns = table_data['columns']

            # Получаем доступ к текущей таблице
            table = doc.tables[table_index]
            
            # Заполняем ячейки таблицы данными из строки Excel
            for i in range(rows):
                for j in range(columns):
                    # Получаем значение ячейки из строки Excel
                    if i * columns + j < len(row_values):
                        cell_value = row_values[i * columns + j]  # Предполагается, что данные из Excel помещаются в таблицу построчно
                    else:
                        cell_value = ""
                    cell = table.cell(i, j)

                    # Проходим по всем параграфам в ячейке и заменяем текстовые поля на значения из Excel
                    for paragraph in cell.paragraphs:
                        for field_index, field_value in enumerate(row_values):
                            field_name = f"{{{field_index}}}"
                            if field_name in paragraph.text:
                                # Проверяем, является ли значение "None" и заменяем на пустую строку
                                if field_value is None:
                                    field_value = ''
                                paragraph.text = paragraph.text.replace(field_name, str(field_value))
                                for run in paragraph.runs:
                                    run.font.name = fname  # устанавливаем шрифт
                                    run.font.size = fsize    # устанавливаем размер текста

        # Создаем имя для нового файла Word
        output_file = f"{output_folder}/{row[0]}.docx"  # Предположим, что первый столбец Excel содержит уникальные идентификаторы

        # Сохраняем новый документ Word
        doc.save(output_file)

fill_word_template(template_path, excel_path, output_folder)
